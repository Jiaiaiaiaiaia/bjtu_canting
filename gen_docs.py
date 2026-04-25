"""生成开发阶段全部交付物 Word 文档（按课程格式规范）。

格式规范要点：
  - 封面：标题一号/二号字居中，作者信息四号字居中
  - 目录：自动生成（插入 TOC 域）
  - 正文：小四宋体，首行缩进2字符，1.5倍行距
  - 标题：黑体，层级递减（三号→小三→四号）
  - 页眉：文档名称，右对齐，小五宋体
  - 页码：正文起始，底部居中
  - 页边距：上下2.54cm，左右3.17cm
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = '/Users/sissi/Downloads/03_软件综合实训/开发阶段交付物'
os.makedirs(OUT, exist_ok=True)

# ── 字号对照（磅值） ──
FONT_1 = 26     # 一号
FONT_2 = 22     # 二号
FONT_3 = 16     # 三号
FONT_S3 = 15    # 小三
FONT_4 = 14     # 四号
FONT_S4 = 12    # 小四
FONT_5 = 10.5   # 五号
FONT_S5 = 9     # 小五

def set_cjk_font(run, name='宋体', size=FONT_S4, bold=False, color=None):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = name
    rpr = run.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    if color:
        run.font.color.rgb = RGBColor(*color)

def setup_styles(doc):
    """配置文档默认样式和标题样式。"""
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(FONT_S4)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(0.74)

    heading_configs = [
        ('Heading 1', '黑体', FONT_3, True),
        ('Heading 2', '黑体', FONT_S3, True),
        ('Heading 3', '黑体', FONT_4, True),
    ]
    for name, font, size, bold in heading_configs:
        s = doc.styles[name]
        s.font.name = font
        s.font.size = Pt(size)
        s.font.bold = bold
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.element.rPr.rFonts.set(qn('w:eastAsia'), font)
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.line_spacing = 1.5
        s.paragraph_format.first_line_indent = None

def setup_page(doc):
    """设置页边距。"""
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def add_header(doc, text):
    """添加页眉（右对齐，小五宋体）。"""
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    set_cjk_font(r, '宋体', FONT_S5)
    # 页眉下划线
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_page_number(doc):
    """添加页码（底部居中）。"""
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('第 ')
    set_cjk_font(r, '宋体', FONT_S5)
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    r2 = p.add_run()
    r2._element.append(fld_begin)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    r3 = p.add_run()
    r3._element.append(instr)
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r4 = p.add_run()
    r4._element.append(fld_end)
    r5 = p.add_run(' 页')
    set_cjk_font(r5, '宋体', FONT_S5)

def add_toc(doc):
    """插入自动目录域（在Word中按F9更新）。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('目  录')
    set_cjk_font(r, '黑体', FONT_2, True)
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(12)

    p2 = doc.add_paragraph()
    p2.paragraph_format.first_line_indent = None
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    r1 = p2.add_run()
    r1._element.append(fld_begin)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r2 = p2.add_run()
    r2._element.append(instr)
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r3 = p2.add_run()
    r3._element.append(fld_end)

    p3 = doc.add_paragraph()
    p3.paragraph_format.first_line_indent = None
    r = p3.add_run('（请在 Word 中右键目录区域 → "更新域" → "更新整个目录"）')
    set_cjk_font(r, '宋体', FONT_5, color=(0x99, 0x99, 0x99))

    doc.add_page_break()

# ── 封面 ──

def make_cover_team(doc, title, subtitle):
    for _ in range(5):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    r = p.add_run('北京交通大学')
    set_cjk_font(r, '黑体', FONT_1, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('软件综合实训')
    set_cjk_font(r, '宋体', FONT_3)

    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_cjk_font(r, '黑体', FONT_2, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(30)
    r = p.add_run(subtitle)
    set_cjk_font(r, '宋体', FONT_4)

    info = [('小组编号', '第 3 小组'), ('指导教师', '杨武杰'), ('提交日期', '2026 年   月   日')]
    t = doc.add_table(rows=len(info), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(info):
        for j, txt in enumerate([k, v]):
            c = t.cell(i, j)
            c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = None
            r = p.add_run(txt)
            set_cjk_font(r, '宋体', FONT_4, j == 0)
        t.cell(i, 0).width = Cm(4)
        t.cell(i, 1).width = Cm(8)

    doc.add_page_break()

def make_cover_personal(doc, title, subtitle, role):
    for _ in range(5):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    r = p.add_run('北京交通大学')
    set_cjk_font(r, '黑体', FONT_1, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('软件综合实训')
    set_cjk_font(r, '宋体', FONT_3)

    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_cjk_font(r, '黑体', FONT_2, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(30)
    r = p.add_run(subtitle)
    set_cjk_font(r, '宋体', FONT_4)

    info = [
        ('学    号', '[         ]'),
        ('姓    名', '[         ]'),
        ('小组编号', '第 3 小组'),
        ('承担角色', role),
        ('指导教师', '杨武杰'),
        ('提交日期', '2026 年   月   日'),
    ]
    t = doc.add_table(rows=len(info), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(info):
        for j, txt in enumerate([k, v]):
            c = t.cell(i, j)
            c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = None
            r = p.add_run(txt)
            set_cjk_font(r, '宋体', FONT_4, j == 0)
        t.cell(i, 0).width = Cm(4)
        t.cell(i, 1).width = Cm(8)

    doc.add_page_break()

# ── 正文辅助 ──

def add_body_para(doc, text, indent=True):
    p = doc.add_paragraph(style='Normal')
    if not indent:
        p.paragraph_format.first_line_indent = None
    r = p.add_run(text)
    set_cjk_font(r, '宋体', FONT_S4)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = None
    r = p.add_run(text)
    set_cjk_font(r, '宋体', FONT_S4)
    return p

def add_number(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = None
    r = p.add_run(text)
    set_cjk_font(r, '宋体', FONT_S4)
    return p

def make_table(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers), style='Table Grid')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        r = p.add_run(h)
        set_cjk_font(r, '黑体', FONT_5, True)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[i+1].cells[j]
            c.text = ''
            p = c.paragraphs[0]
            p.paragraph_format.first_line_indent = None
            r = p.add_run(str(val))
            set_cjk_font(r, '宋体', FONT_5)
    return t

def new_doc():
    doc = Document()
    setup_styles(doc)
    setup_page(doc)
    return doc

# ═══════════════════════════════════════════════════════
# 文档 1: 开发阶段小组沟通交流记录
# ═══════════════════════════════════════════════════════
def gen_communication_record():
    doc = new_doc()
    make_cover_team(doc, '开发阶段小组沟通交流记录', '北京交通大学就餐仿真系统')
    add_toc(doc)
    add_header(doc, '03小组_开发阶段小组沟通交流记录')
    add_page_number(doc)

    doc.add_heading('一、概述', level=1)
    add_body_para(doc, '本文档记录了第3小组在开发阶段（第5-9周）的主要沟通交流活动。开发阶段的核心目标是完成各模块的编码实现与单元测试，并为后续集成联调做好准备。小组共进行了5次正式沟通交流，形式包括线下面对面讨论和线上微信群沟通。')

    meetings = [
        {
            'num': 1, 'date': '2026年4月5日（第5周）', 'place': '教九楼机房',
            'attendees': '全体组员（3人）', 'duration': '45分钟',
            'topic': '开发阶段启动与任务确认',
            'content': [
                '回顾立项阶段成果，确认设计规格说明书中各模块接口定义；',
                '讨论并确认每人负责的模块范围及开发优先级；',
                '约定代码提交规范：每完成一个功能点提交一次，commit message 使用中文描述；',
                '确定每周至少沟通一次进度，遇到阻塞问题及时在微信群反馈。',
            ],
            'result': '明确三人分工：后端开发负责仿真引擎与API，前端开发负责界面与可视化，配置与分析负责数据库与统计指标。制定了第5-6周完成核心模块、第7-8周完成单元测试的里程碑。',
        },
        {
            'num': 2, 'date': '2026年4月12日（第6周）', 'place': '线上微信群',
            'attendees': '全体组员（3人）', 'duration': '30分钟',
            'topic': '核心模块开发进度同步',
            'content': [
                '后端开发汇报：仿真驱动模块（engine.py）核心逻辑已完成，离散事件队列运行正常；排队仿真模块（queue_sim.py）和就餐仿真模块（dining_sim.py）已完成编码；',
                '前端开发汇报：参数配置页和仿真运行页HTML结构已完成，Canvas食堂布局绘制基本成型；',
                '配置与分析汇报：SQLite数据库建表完成，配置参数校验逻辑已实现；',
                '讨论了前后端数据交互格式，统一了/api/simulation/step接口返回的JSON字段命名。',
            ],
            'result': '确认核心模块开发进度符合预期。前端对后端返回的students数组格式提出了调整建议，后端同意增加position和position_detail字段以支持Canvas动画渲染。',
        },
        {
            'num': 3, 'date': '2026年4月19日（第7周）', 'place': '教九楼机房',
            'attendees': '全体组员（3人）', 'duration': '50分钟',
            'topic': '前后端联调与界面优化讨论',
            'content': [
                '进行了首次前后端联调测试，验证了参数配置→仿真启动→逐步推进→结束仿真的完整流程；',
                '发现问题：Canvas绘制的图例与窗口颜色不统一，速度控制使用下拉框操作不够直观；',
                '讨论了界面视觉风格优化方案，决定采用北交大红（#b91c1c）作为主色调，使用学术简约设计风格；',
                '决定将速度控制从下拉框改为滑块（range input），并在仿真信息面板增加"平均等待时间"实时指标。',
            ],
            'result': '完成首次联调，确认核心流程无阻塞性问题。后端需新增avg_waiting_time字段到step接口返回值中；前端需完成CSS重写和JS适配工作。',
        },
        {
            'num': 4, 'date': '2026年4月26日（第8周）', 'place': '线上微信群',
            'attendees': '全体组员（3人）', 'duration': '35分钟',
            'topic': '单元测试与文档编写',
            'content': [
                '后端开发汇报：39个单元测试用例全部通过，覆盖仿真引擎生命周期、排队模块、就餐模块和API接口；',
                '前端开发汇报：CSS重写完成，JS速度滑块和平均等待时间显示已适配，ECharts图表颜色统一为主题色；',
                '配置与分析汇报：数据库读写功能测试通过，统计指标计算逻辑验证正确；',
                '讨论了单元测试报告和开发阶段实训报告的撰写分工。',
            ],
            'result': '各模块开发和测试均已完成。约定第9周前完成所有交付文档的撰写，并进行交叉审阅。',
        },
        {
            'num': 5, 'date': '2026年5月2日（第9周）', 'place': '线上微信群',
            'attendees': '全体组员（3人）', 'duration': '25分钟',
            'topic': '交付物确认与提交前检查',
            'content': [
                '逐项检查开发阶段交付物清单，确认团队文档和个人文档均已完成；',
                '交叉审阅了各自的单元测试报告，修正了部分格式问题；',
                '确认源代码打包方式和命名规范符合课程要求。',
            ],
            'result': '所有交付物确认就绪，约定5月3日前完成最终提交。',
        },
    ]

    doc.add_heading('二、沟通交流记录', level=1)
    for m in meetings:
        doc.add_heading(f'2.{m["num"]} 第{m["num"]}次沟通交流', level=2)
        info_table = [
            ['时间', m['date']],
            ['地点', m['place']],
            ['参加人员', m['attendees']],
            ['时长', m['duration']],
            ['主题', m['topic']],
        ]
        t = doc.add_table(rows=5, cols=2, style='Table Grid')
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (k, v) in enumerate(info_table):
            for j, txt in enumerate([k, v]):
                c = t.cell(i, j)
                c.text = ''
                p = c.paragraphs[0]
                p.paragraph_format.first_line_indent = None
                r = p.add_run(txt)
                set_cjk_font(r, '宋体' if j == 1 else '黑体', FONT_5, j == 0)
            t.cell(i, 0).width = Cm(3)

        doc.add_paragraph()
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
        r = p.add_run('交流内容：')
        set_cjk_font(r, '黑体', FONT_S4, True)

        for item in m['content']:
            add_bullet(doc, item)

        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
        r = p.add_run('议定结果：')
        set_cjk_font(r, '黑体', FONT_S4, True)
        add_body_para(doc, m['result'])
        doc.add_paragraph()

    doc.save(os.path.join(OUT, '03小组_开发阶段小组沟通交流记录.docx'))
    print('  ✓ 沟通交流记录')


# ═══════════════════════════════════════════════════════
# 文档 2: 小组开发任务划分说明
# ═══════════════════════════════════════════════════════
def gen_task_division():
    doc = new_doc()
    make_cover_team(doc, '小组开发任务划分说明', '北京交通大学就餐仿真系统')
    add_toc(doc)
    add_header(doc, '03小组_小组开发任务划分说明')
    add_page_number(doc)

    doc.add_heading('一、任务划分原则', level=1)
    add_body_para(doc, '本小组按照设计规格说明书中的模块划分进行任务分配，每位成员独立负责一个或多个完整模块的开发与测试，确保各模块边界清晰、接口明确、可独立验证。')

    doc.add_heading('二、成员角色与负责模块', level=1)
    make_table(doc,
        ['成员', '学号', '角色', '负责模块', '主要职责'],
        [
            ['[姓名1]', '[学号1]', '后端开发', '仿真驱动模块\n排队仿真模块\n就餐仿真模块\nAPI接口模块', '离散事件仿真引擎核心代码\nRESTful API接口开发\n后端单元测试'],
            ['[姓名2]', '[学号2]', '前端开发', '可视化模块\n前端交互', 'HTML页面结构\nCSS样式设计\nCanvas食堂布局绘制\nECharts图表展示\nJS交互逻辑'],
            ['[姓名3]', '[学号3]', '配置与分析', '初始配置模块\n数据分析模块\n数据存储', '参数配置校验\nSQLite数据库设计与读写\n统计指标计算\n历史记录管理'],
        ])

    doc.add_heading('三、各模块详细任务分解', level=1)

    doc.add_heading('3.1 后端开发任务', level=2)
    make_table(doc,
        ['序号', '任务', '涉及文件', '预计工期', '完成状态'],
        [
            ['1', '实现离散事件驱动仿真引擎', 'simulation/engine.py', '第5-6周', '已完成'],
            ['2', '实现排队仿真模块', 'simulation/queue_sim.py', '第5周', '已完成'],
            ['3', '实现就餐仿真模块', 'simulation/dining_sim.py', '第5周', '已完成'],
            ['4', '实现RESTful API接口', 'api/routes.py', '第6周', '已完成'],
            ['5', '新增实时平均等待时间指标', 'simulation/engine.py', '第7周', '已完成'],
            ['6', '编写后端单元测试（39个用例）', 'tests/*.py', '第7-8周', '已完成'],
        ])

    doc.add_heading('3.2 前端开发任务', level=2)
    make_table(doc,
        ['序号', '任务', '涉及文件', '预计工期', '完成状态'],
        [
            ['1', '实现参数配置页', 'index.html, main.js', '第5周', '已完成'],
            ['2', '实现Canvas食堂布局绘制', 'main.js', '第5-6周', '已完成'],
            ['3', '实现学生级动画渲染', 'main.js', '第6周', '已完成'],
            ['4', '实现ECharts图表展示', 'main.js', '第6-7周', '已完成'],
            ['5', '实现历史记录页', 'main.js', '第7周', '已完成'],
            ['6', 'CSS重写为学术简约设计风格', 'style.css', '第7-8周', '已完成'],
            ['7', 'JS适配速度滑块与颜色统一', 'main.js', '第8周', '已完成'],
        ])

    doc.add_heading('3.3 配置与分析任务', level=2)
    make_table(doc,
        ['序号', '任务', '涉及文件', '预计工期', '完成状态'],
        [
            ['1', 'SQLite数据库表设计与建表', 'api/routes.py', '第5周', '已完成'],
            ['2', '参数配置校验逻辑实现', 'api/routes.py', '第5周', '已完成'],
            ['3', '仿真快照批量写入与读取', 'api/routes.py', '第6周', '已完成'],
            ['4', '统计指标计算模块', 'simulation/engine.py', '第6-7周', '已完成'],
            ['5', '历史记录查询接口', 'api/routes.py', '第7周', '已完成'],
            ['6', '数据分析模块测试与验证', 'tests/*.py', '第8周', '已完成'],
        ])

    doc.add_heading('四、协作方式', level=1)
    add_number(doc, '版本管理：使用Git进行版本控制，各成员在各自模块目录下独立开发；')
    add_number(doc, '沟通方式：每周至少一次线上/线下沟通，同步开发进度和协调接口变更；')
    add_number(doc, '接口约定：前后端通过RESTful API进行数据交互，接口定义以设计规格说明书为准，变更需经小组讨论确认；')
    add_number(doc, '测试策略：各模块开发完成后先进行单元测试，再进行前后端联调验证。')

    doc.save(os.path.join(OUT, '03小组_小组开发任务划分说明.docx'))
    print('  ✓ 任务划分说明')


# ═══════════════════════════════════════════════════════
# 文档 3-5: 单元测试报告 × 3
# ═══════════════════════════════════════════════════════
def gen_test_report(role, role_label, tc):
    doc = new_doc()
    make_cover_personal(doc, '单元测试报告', '北京交通大学就餐仿真系统', role_label)
    add_toc(doc)
    add_header(doc, f'软件综合实训_[学号]_[姓名]_单元测试报告')
    add_page_number(doc)

    doc.add_heading('一、测试概述', level=1)
    add_body_para(doc, tc['overview'])

    doc.add_heading('二、测试目标', level=1)
    add_body_para(doc, tc['goal'])

    doc.add_heading('三、测试环境', level=1)
    make_table(doc, ['项目', '内容'], [
        ['操作系统', 'macOS 14 / Windows 11'],
        ['Python版本', 'Python 3.12'],
        ['测试框架', 'pytest 9.0'],
        ['开发工具', 'PyCharm / VS Code'],
        ['浏览器', 'Chrome 124（前端测试）'],
    ])

    doc.add_heading('四、测试用例设计', level=1)
    add_body_para(doc, tc['design_intro'])
    for sec in tc['test_sections']:
        doc.add_heading(sec['title'], level=2)
        add_body_para(doc, sec['desc'])
        make_table(doc, ['用例编号', '测试内容', '输入/前置条件', '预期结果', '实际结果'], sec['cases'])

    doc.add_heading('五、测试执行与结果', level=1)
    add_body_para(doc, tc['execution'])

    doc.add_heading('六、问题与解决', level=1)
    if tc.get('issues'):
        make_table(doc, ['序号', '发现的问题', '问题原因', '解决方法'], tc['issues'])
    else:
        add_body_para(doc, '本轮测试未发现严重缺陷。')

    doc.add_heading('七、测试结论', level=1)
    add_body_para(doc, tc['conclusion'])

    fname = f'软件综合实训_[学号]_[姓名]_单元测试报告_{role}.docx'
    doc.save(os.path.join(OUT, fname))
    print(f'  ✓ 单元测试报告 ({role_label})')


def gen_all_test_reports():
    gen_test_report('后端开发', '后端开发', {
        'overview': '本报告为后端开发成员负责的仿真引擎模块、排队仿真模块、就餐仿真模块及API接口模块的单元测试报告。测试使用pytest框架编写，共计39个测试用例，覆盖了模块的核心逻辑、边界条件和不变量验证。',
        'goal': '验证离散事件驱动仿真引擎在各种配置参数下的正确性，包括：学生到达计数的单调性、座位剩余时间的动态递减、仿真结束时的不变量（到达数等于完成数）、统计指标的合理性，以及API接口的参数校验与状态管理。',
        'design_intro': '测试用例按模块组织，采用pytest fixture提供标准配置和引擎实例。使用固定随机种子（rng_seed=42）保证测试结果可复现。测试重点验证系统不变量而非具体数值，以适应随机仿真的特性。',
        'test_sections': [
            {'title': '4.1 仿真引擎模块测试', 'desc': '测试SimulationEngine类的完整生命周期，包括初始状态、启动保护、到达计数、座位剩余时间、平均等待时间、运行完成不变量和统计指标。',
             'cases': [
                ['E-01', '初始计数器为零', '创建引擎实例，不调用start()', 'total_arrived=0, total_served=0', '通过'],
                ['E-02', '窗口和座位正确初始化', 'window_count=3, seat_count=30', '创建3个窗口、30个座位', '通过'],
                ['E-03', 'start()幂等性', '连续调用start()两次', '事件队列和学生数不变', '通过'],
                ['E-04', 'start()后arrived仍为0', '调用start()后立即检查', 'total_arrived=0, students非空', '通过'],
                ['E-05', 'arrived单调递增', '运行100步', '每步arrived >= 上一步', '通过'],
                ['E-06', 'arrived不超过预生成学生数', '运行至结束', 'total_arrived <= len(students)', '通过'],
                ['E-07', '座位remaining_time递减', '观察同一座位连续状态', '后续remaining_time < 初始值', '通过'],
                ['E-08', '空座remaining_time为0', '检查空座位状态', 'remaining_time == 0', '通过'],
                ['E-09', 'avg_waiting_time存在且非负', '运行51步后检查', '字段存在，类型为数值，>= 0', '通过'],
                ['E-10', '结束时arrived == served', '运行至is_ended', '两者相等且大于0', '通过'],
                ['E-11', '结束时无残留状态', '运行至is_ended', '座位空、窗口空、等位队列空', '通过'],
                ['E-12', '统计指标基本不变量', '运行至结束后调用get_statistics()', '座位利用率0-100%，窗口served之和等于total_served', '通过'],
                ['E-13', '时间线数据非空', '运行至结束后调用get_statistics()', 'queue_timeline有数据，长度>=5', '通过'],
            ]},
            {'title': '4.2 排队仿真模块测试', 'desc': '测试Window类、最短队列分配策略和打饭时长采样函数。',
             'cases': [
                ['Q-01', 'Window初始状态', '创建Window实例', 'queue为空，current_serving为None', '通过'],
                ['Q-02', 'queue_load含正在服务', '设置current_serving后检查', 'queue_load() = queue长度 + 1', '通过'],
                ['Q-03', 'pick_shortest_window', '3个窗口不同负载', '返回负载最小的窗口', '通过'],
                ['Q-04', 'serving计入负载', '窗口正在服务且队列为空', '该窗口负载=1', '通过'],
                ['Q-05', 'serve_time下限1秒', '采样1000次', '所有值 >= 1.0', '通过'],
                ['Q-06', 'serve_time均值接近配置', '采样10000次，avg=30', '均值在27-33之间', '通过'],
                ['Q-07', 'serve_time无负值', '采样1000次', '所有值 > 0', '通过'],
            ]},
            {'title': '4.3 就餐仿真模块测试', 'desc': '测试Seat类、就近座位分配策略和就餐时长采样函数。',
             'cases': [
                ['D-01', 'Seat初始状态', '创建Seat实例', "status='empty', student=None", '通过'],
                ['D-02', '全部占满返回None', '所有座位occupied', 'pick_nearest_seat返回None', '通过'],
                ['D-03', '就近分配空闲座位', '部分座位空闲', '返回距离window_id最近的空座', '通过'],
                ['D-04', '跳过已占座位', '目标位置已占用', '返回下一个最近的空座', '通过'],
                ['D-05', 'eat_time下限60秒', '采样1000次', '所有值 >= 60.0', '通过'],
                ['D-06', 'eat_time均值接近配置', '采样10000次，avg=15min', '均值在13-17分钟之间', '通过'],
            ]},
            {'title': '4.4 API接口测试', 'desc': '测试RESTful API的参数校验、仿真控制和数据查询接口。使用Flask测试客户端进行请求模拟。',
             'cases': [
                ['A-01', '合法配置接受', 'POST /api/config 完整参数', '200，返回config_id', '通过'],
                ['A-02', '缺少字段拒绝', '省略seat_count', '400，错误信息', '通过'],
                ['A-03', '负数值拒绝', 'window_count=-1', '400，错误信息', '通过'],
                ['A-04', '首次start正常', 'POST /api/simulation/start', '200，status=running', '通过'],
                ['A-05', '重复start不污染', '连续两次start', '第二次返回already_started', '通过'],
                ['A-06', '未配置时start返回400', '不提交config直接start', '400', '通过'],
                ['A-07', '未运行时step返回400', '不start直接step', '400', '通过'],
                ['A-08', 'step返回完整状态', '正常运行后GET step', '包含current_time等字段', '通过'],
                ['A-09', 'finish返回完整统计', 'POST finish', '包含total_arrived等', '通过'],
                ['A-10', '未配置时finish返回400', '不配置直接finish', '400', '通过'],
                ['A-11', '历史配置列表', 'GET /api/history/configs', '返回配置数组', '通过'],
                ['A-12', '历史快照查询', 'GET /api/history?config_id=N', '返回快照数组', '通过'],
                ['A-13', 'reset清除引擎', 'POST reset后检查', '引擎为None', '通过'],
            ]},
        ],
        'execution': '使用命令 python3 -m pytest tests/ -v 执行全部测试，39个用例全部通过，执行时间约0.4秒。测试覆盖了仿真引擎的完整生命周期、各子模块的核心逻辑和API接口的正常与异常路径。',
        'issues': [
            ['1', '座位remaining_time为静态值', '入座时记录固定就餐时长，未随current_time递减', '修改为eat_end_time - current_time动态计算'],
            ['2', '平均等待时间口径不一致', '_build_state与get_statistics使用不同学生集合', '明确区分实时口径与终局口径，仿真结束时收敛一致'],
        ],
        'conclusion': '经过39个测试用例的全面验证，后端仿真引擎模块、排队仿真模块、就餐仿真模块和API接口模块均运行正确。所有核心不变量（到达数等于完成数、座位利用率范围合理、统计指标一致性）均得到满足。模块可以进入集成联调阶段。',
    })

    gen_test_report('前端开发', '前端开发', {
        'overview': '本报告为前端开发成员负责的可视化模块的测试报告。由于前端代码主要涉及Canvas绑制、DOM交互和ECharts图表渲染，采用人工操作测试与浏览器控制台验证相结合的方式进行测试。',
        'goal': '验证前端四个页面（参数配置、仿真运行、数据分析、历史记录）的交互逻辑正确性、界面渲染效果和响应式布局适配。',
        'design_intro': '测试用例按页面组织，每个页面覆盖正常操作流程、边界条件和异常输入处理。测试在Chrome浏览器中进行，使用开发者工具监控网络请求和控制台输出。',
        'test_sections': [
            {'title': '4.1 参数配置页测试', 'desc': '测试表单输入、参数校验、默认值恢复和提交功能。',
             'cases': [
                ['F-01', '默认值正确显示', '打开配置页', '窗口数=6,座位=200等', '通过'],
                ['F-02', '恢复默认值按钮', '修改参数后点击恢复', '所有输入恢复为默认值', '通过'],
                ['F-03', '提交成功跳转', '填写合法参数点击开始', '跳转到仿真运行页', '通过'],
                ['F-04', 'HTML5 min/max约束', '输入超范围值', '浏览器原生校验阻止提交', '通过'],
                ['F-05', '后端无响应时提示', '关闭后端后点击开始', '弹出连接失败提示', '通过'],
            ]},
            {'title': '4.2 仿真运行页测试', 'desc': '测试Canvas绘制、信息面板更新、速度控制和仿真流程控制。',
             'cases': [
                ['F-06', '7个信息卡片显示', '进入仿真运行页', '显示全部7项指标', '通过'],
                ['F-07', '平均等待实时更新', '运行仿真观察', '数值随step更新，非负', '通过'],
                ['F-08', '速度滑块响应', '拖动速度滑块', '标签显示×1/×2/×5/×10', '通过'],
                ['F-09', '暂停/继续切换', '点击暂停再继续', '按钮文字正确切换', '通过'],
                ['F-10', 'Canvas窗口颜色', '观察Canvas绘制', '占用窗口红色，空闲灰色', '通过'],
                ['F-11', '学生圆点动画', '观察Canvas动画', '帧间插值平滑移动', '通过'],
                ['F-12', 'HTML图例栏', '观察Canvas上方', '显示8个图例项', '通过'],
                ['F-13', '结束仿真跳转', '点击结束仿真', '跳转到数据分析页', '通过'],
                ['F-14', '自然结束自动跳转', '等待仿真跑完', '自动跳转到数据分析页', '通过'],
            ]},
            {'title': '4.3 数据分析页测试', 'desc': '测试统计卡片和ECharts图表的数据展示。',
             'cases': [
                ['F-15', '6个统计卡片', '进入分析页', '显示全部6项统计数据', '通过'],
                ['F-16', '柱状图窗口数据', '检查柱状图', '窗口数量与配置一致', '通过'],
                ['F-17', '折线图趋势合理', '检查排队变化图', '呈现先升后降趋势', '通过'],
                ['F-18', '重新仿真按钮', '点击重新仿真', '跳转回配置页', '通过'],
                ['F-19', '图表自适应', '调整窗口大小', '图表跟随调整', '通过'],
            ]},
            {'title': '4.4 历史记录页测试', 'desc': '测试历史记录表格和详情图表。',
             'cases': [
                ['F-20', '刷新列表', '点击刷新列表', '表格显示历史记录', '通过'],
                ['F-21', '空记录提示', '清空数据库后查看', '显示"暂无历史记录"', '通过'],
                ['F-22', '点击行显示详情', '点击某条记录', '下方显示时间曲线图', '通过'],
                ['F-23', '行高亮与切换', '连续点击不同行', '选中行高亮，上一行取消', '通过'],
                ['F-24', '斑马条纹', '观察表格行', '奇数行浅灰背景', '通过'],
            ]},
        ],
        'execution': '在Chrome 124浏览器中完成全部24个测试用例的人工验证，同时使用开发者工具监控网络请求和JavaScript控制台，确认无报错信息。',
        'issues': [
            ['1', 'Canvas图例颜色不统一', '使用旧蓝色而非主题红色', '统一更新颜色，移除Canvas图例改用HTML图例栏'],
            ['2', '速度下拉框不直观', '需点开再选择', '改用range滑块即时生效'],
        ],
        'conclusion': '经过24个测试用例的全面验证，前端可视化模块四个页面均功能正常、界面美观、交互流畅。Canvas动画渲染效果良好，ECharts图表数据展示准确。CSS响应式布局在不同窗口宽度下均能正常适配。',
    })

    gen_test_report('配置与分析', '配置与分析', {
        'overview': '本报告为配置与分析成员负责的初始配置模块、数据分析模块和数据存储功能的单元测试报告。测试采用pytest框架和Flask测试客户端，验证参数校验、数据库读写和统计指标计算的正确性。',
        'goal': '验证参数配置校验的完备性、SQLite数据库读写的正确性和统计指标计算的准确性。',
        'design_intro': '测试重点覆盖配置校验的正常与异常路径、数据库CRUD操作的完整性、以及统计指标计算结果与数学期望的一致性。',
        'test_sections': [
            {'title': '4.1 参数配置校验测试', 'desc': '测试_validate_config函数对各种输入的校验结果。',
             'cases': [
                ['C-01', '完整合法配置通过', '6个字段均为正值', '返回None', '通过'],
                ['C-02', '缺少必填字段', '省略seat_count', '返回错误信息', '通过'],
                ['C-03', '窗口数为0', 'window_count=0', '返回错误信息', '通过'],
                ['C-04', '负数到达率', 'arrival_rate=-5', '返回错误信息', '通过'],
                ['C-05', '字符串类型参数', "window_count='abc'", '返回类型错误', '通过'],
            ]},
            {'title': '4.2 数据库读写测试', 'desc': '测试配置写入、快照批量写入和历史查询功能。',
             'cases': [
                ['C-06', '配置写入并返回ID', 'POST /api/config', '返回递增config_id', '通过'],
                ['C-07', '快照批量写入', '运行50步后触发flush', '数据库有对应记录', '通过'],
                ['C-08', 'finish后快照完整', 'POST finish', '快照数等于总事件数', '通过'],
                ['C-09', '历史配置查询', 'GET /api/history/configs', '返回含snapshot_count的列表', '通过'],
                ['C-10', '按config_id查快照', 'GET /api/history?config_id=1', '仅返回对应快照', '通过'],
                ['C-11', 'reset后可重新配置', 'POST reset后重新config', '新config_id递增', '通过'],
            ]},
            {'title': '4.3 统计指标计算测试', 'desc': '测试get_statistics()返回的各项统计指标。',
             'cases': [
                ['C-12', '总到达数一致', '运行至结束', 'stats与engine一致', '通过'],
                ['C-13', '总完成数一致', '运行至结束', 'stats与engine一致', '通过'],
                ['C-14', '座位利用率范围', '运行至结束', '0-100%', '通过'],
                ['C-15', '窗口served之和', '运行至结束', '等于total_served', '通过'],
                ['C-16', '平均等待非负', '运行至结束', '>= 0', '通过'],
                ['C-17', '平均就餐合理', 'avg_eat_time=5min', '>= 60秒', '通过'],
                ['C-18', '排队时间线覆盖', '运行5分钟', '至少5个采样点', '通过'],
            ]},
        ],
        'execution': '使用pytest执行相关测试用例，所有用例均通过。统计指标通过固定随机种子和已知期望值进行交叉验证。',
        'issues': [['1', '座位利用率超过100%', '有效时间仅取配置时长', '改为max(配置时长, 实际结束时间)']],
        'conclusion': '经过18个测试用例的验证，参数配置校验完备、数据库读写可靠、统计指标计算准确。所有核心不变量均满足，数据模块可以进入集成联调阶段。',
    })


# ═══════════════════════════════════════════════════════
# 文档 6-8: 开发阶段实训报告 × 3
# ═══════════════════════════════════════════════════════
def gen_dev_report(role, role_label, c):
    doc = new_doc()
    make_cover_personal(doc, '开发阶段个人实训报告', '北京交通大学就餐仿真系统', role_label)
    add_toc(doc)
    add_header(doc, f'软件综合实训_[学号]_[姓名]_开发阶段实训报告')
    add_page_number(doc)

    doc.add_heading('一、实训概述', level=1)
    doc.add_heading('1.1 阶段背景', level=2)
    add_body_para(doc, c['background'])
    doc.add_heading('1.2 阶段时间', level=2)
    add_body_para(doc, '2026年4月5日 —— 2026年5月3日（第5-9周）')

    doc.add_heading('二、遇到的问题及解决方式', level=1)
    add_body_para(doc, '在本阶段实训过程中，本人遇到了以下主要问题，并尝试采用了相应的方法加以解决：')
    make_table(doc, ['序号', '遇到的问题', '问题详细描述', '解决方法'], c['problems'])

    doc.add_heading('三、个人承担任务与完成情况', level=1)
    doc.add_heading('3.1 任务承担总览', level=2)
    add_body_para(doc, c['task_overview'])
    make_table(doc, ['序号', '承担的任务', '完成情况', '遇到的问题', '解决方式'], c['tasks'])

    doc.add_heading('3.2 任务完成详细说明', level=2)
    add_body_para(doc, '在上述任务中，以下两项任务的完成过程令本人印象最深：')
    for i, detail in enumerate(c['task_details'], 1):
        add_body_para(doc, f'{i}. {detail}')

    doc.add_heading('四、阶段学习与收获', level=1)
    doc.add_heading('4.1 技术收获', level=2)
    add_body_para(doc, '通过本阶段的学习，本人在以下技术方面获得了显著提升：')
    for item in c['tech_gains']:
        add_bullet(doc, item)

    doc.add_heading('4.2 过程态度与团队协作感想', level=2)
    add_body_para(doc, c['reflection'])

    doc.add_heading('五、下一阶段计划', level=1)
    add_body_para(doc, '在接下来的集成阶段（第10-12周），本人计划完成以下目标：')
    for item in c['next_plans']:
        add_number(doc, item)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = None
    r = p.add_run('报告人：[姓名]        日期：2026年   月   日')
    set_cjk_font(r, '宋体', FONT_S4)

    fname = f'软件综合实训_[学号]_[姓名]_开发阶段实训报告_{role}.docx'
    doc.save(os.path.join(OUT, fname))
    print(f'  ✓ 开发阶段实训报告 ({role_label})')


def gen_all_dev_reports():
    gen_dev_report('后端开发', '后端开发', {
        'background': '本报告为软件综合实训课程开发阶段的个人总结报告。开发阶段的主要任务是根据立项阶段的设计规格说明书，完成各模块的编码实现与单元测试。在本阶段中，我继续承担后端开发工作，负责仿真引擎核心代码的实现、RESTful API接口开发以及后端模块的单元测试。',
        'problems': [
            ['1', '离散事件驱动模型的实现复杂度', '事件之间的依赖关系较为复杂，特别是arrival→service_end→eat_end的事件链中，每个事件处理完后需要正确触发后续事件，且需要维护全局状态的一致性。', '采用最小堆（heapq）实现事件优先队列，每个事件携带时间戳和序号保证有序性；将事件处理逻辑拆分为三个独立方法，降低耦合度。'],
            ['2', '实时平均等待时间的计算口径问题', '实时统计口径（已开始服务的学生）与最终统计口径（已完成就餐的学生）的计算基数不同，导致中途数值与最终结果不一致。', '明确两种口径的适用场景：实时面板使用"已开始服务"口径，最终统计使用"已完成就餐"口径，两者在仿真结束后自然收敛一致。'],
            ['3', '单元测试中的随机性控制', '仿真涉及大量随机数，直接测试结果不可复现。', '增加rng_seed参数，测试时使用固定种子；设计测试用例重点验证不变量而非具体数值。'],
        ],
        'task_overview': '本人在本阶段的角色定位为后端开发责任人，主要承担仿真引擎实现、API接口开发和后端单元测试。具体任务如下表所示：',
        'tasks': [
            ['1', '实现离散事件驱动仿真引擎', '已完成', '事件链依赖复杂', '拆分事件处理方法'],
            ['2', '实现排队仿真模块', '已完成', '无', '—'],
            ['3', '实现就餐仿真模块', '已完成', '就近座位分配算法', '按窗口位置映射座位区域'],
            ['4', '实现RESTful API接口', '已完成', '快照写入性能瓶颈', '缓冲区批量写入'],
            ['5', '新增实时avg_waiting_time', '已完成', '计算口径差异', '区分两种口径'],
            ['6', '编写后端单元测试（39个）', '已完成', '随机性不可复现', '固定随机种子'],
        ],
        'task_details': [
            '离散事件驱动仿真引擎的实现：这是本阶段最核心也最具挑战性的任务。核心难点在于事件的正确调度——每个学生会依次经历到达、开始打饭、打饭完成、找座、开始就餐、就餐完成六个阶段，每个阶段转换都需要正确更新全局状态并可能触发新事件。最终采用最小堆加序号的双重排序策略解决了同时刻事件的处理顺序问题。',
            '后端单元测试的设计与实现：最大的收获是学会了如何测试包含随机性的系统——不是去验证具体的数值，而是验证系统的不变量（如仿真结束时所有学生都已离场、到达数等于完成数等）。这种"测试性质而非数值"的思路对我理解测试驱动开发帮助很大。',
        ],
        'tech_gains': [
            '离散事件仿真编程：掌握了基于事件驱动的仿真引擎设计与实现，理解了事件队列、状态快照和仿真循环的核心概念；',
            'Python heapq模块：熟练使用最小堆实现优先队列，理解了自定义比较器的应用；',
            'Flask RESTful API开发：完成了一整套API接口，掌握了Blueprint路由组织、请求参数校验和JSON序列化；',
            'pytest单元测试：学会了fixture、参数化测试和测试组织结构，以及如何设计针对随机系统的不变量测试；',
            'SQLite数据库操作：掌握了Python sqlite3模块的基本用法，包括建表、批量插入和聚合查询。',
        ],
        'reflection': '开发阶段是整个课程中编码量最大的阶段。最深刻的体会有两点：一是接口设计的重要性——在与前端成员联调时，发现立项阶段定义的接口有些字段命名不够直观，经过讨论后增加了position和position_detail字段，大大方便了前端渲染。二是测试先行的价值——在实现avg_waiting_time功能时先写了测试用例再写实现代码，这种TDD方式帮助我提前明确了计算口径，避免了返工。',
        'next_plans': [
            '参与前后端集成联调，确保仿真全流程端到端无阻塞；',
            '配合团队完成接口联调测试用例的编写与执行；',
            '根据联调中发现的问题对后端代码进行修复和优化；',
            '配合撰写集成阶段的联调测试报告。',
        ],
    })

    gen_dev_report('前端开发', '前端开发', {
        'background': '本报告为软件综合实训课程开发阶段的个人总结报告。本人在本阶段承担前端开发工作，负责系统的界面设计与实现，包括HTML页面结构、CSS样式设计、Canvas食堂布局绘制、ECharts图表展示以及JavaScript交互逻辑的编码实现。',
        'problems': [
            ['1', 'Canvas食堂布局的坐标计算', '窗口、队列、座位、学生圆点等多种元素需要在同一画布上分层绘制，坐标计算和渲染顺序的管理较为复杂，特别是在窗口数量可配置的情况下需要动态计算位置。', '将绘制逻辑拆分为drawBackground、drawWindows、drawSeats、drawStudentDots等独立函数，坐标使用相对比例方式动态调整。'],
            ['2', '学生动画的流畅性问题', '直接根据后端返回的位置数据绘制学生圆点会导致学生在不同位置间"瞬移"，视觉效果生硬。', '引入帧间插值（lerp）机制，维护studentPrev状态记录每个学生上一帧位置，按插值系数0.4平滑过渡到目标位置。'],
            ['3', 'CSS视觉风格重写工作量大', '原有CSS样式基础较简单，改为学术简约设计风格需要全面重写，涉及配色方案、圆角、间距、响应式布局等。', '使用CSS自定义属性定义设计令牌，将颜色、圆角等统一管理；采用Grid和Flexbox布局实现响应式适配。'],
        ],
        'task_overview': '本人在本阶段的角色定位为前端开发，负责系统全部界面的设计与实现。具体任务如下表所示：',
        'tasks': [
            ['1', '实现参数配置页', '已完成', 'min/max范围设定', '参考后端校验规则统一'],
            ['2', '实现Canvas食堂布局', '已完成', '坐标动态计算', '按比例分层绘制'],
            ['3', '实现学生级动画渲染', '已完成', '动画流畅性', '引入lerp插值'],
            ['4', '实现ECharts图表展示', '已完成', '图表自适应', '绑定resize事件'],
            ['5', '实现历史记录页', '已完成', '无', '—'],
            ['6', 'CSS全面重写', '已完成', '工作量大', 'CSS自定义属性统一管理'],
            ['7', 'JS适配速度滑块等', '已完成', 'DOM元素变更', '同步更新事件绑定'],
        ],
        'task_details': [
            'Canvas食堂布局与学生动画渲染：在Canvas上实现了食堂平面图的分层绘制，包括窗口（带服务状态颜色标识）、座位矩阵（带就餐热力效果）和学生移动动画。最大的挑战是学生动画的流畅性——最初学生圆点在位置变化时会"跳跃"，引入帧间插值后效果大幅改善。',
            'CSS学术简约设计系统的实现：使用CSS自定义属性定义了完整的设计令牌（颜色、圆角、字体栈），基于这些变量实现各组件样式。最满意的部分是响应式布局——通过两个媒体查询断点实现了从7列到3列再到单列的优雅降级。',
        ],
        'tech_gains': [
            'Canvas 2D绘图API：熟练掌握了fillRect、arc、fillText等绘制方法，理解了分层渲染机制；',
            'CSS自定义属性与设计系统：学会了使用:root变量构建可维护的设计令牌体系；',
            'ECharts图表库：掌握了柱状图、饼图、折线图、面积图的配置方法和生命周期管理；',
            'JavaScript异步编程：通过fetch API和async/await掌握了前后端数据交互模式；',
            '帧间插值动画：理解了线性插值（lerp）在实时渲染中的应用。',
        ],
        'reflection': '开发阶段中最大的收获是学会了"先原型后优化"的开发节奏。初期先确保功能正确，再追求视觉效果。另一个收获是与后端成员的接口协调——通过约定清晰的JSON字段命名，前后端可以高效并行开发。',
        'next_plans': [
            '参与前后端集成联调测试，验证四个页面的端到端功能；',
            '根据联调测试结果修复前端显示问题；',
            '对界面进行细节优化和浏览器兼容性测试；',
            '配合撰写集成阶段的联调测试报告。',
        ],
    })

    gen_dev_report('配置与分析', '配置与分析', {
        'background': '本报告为软件综合实训课程开发阶段的个人总结报告。本人在本阶段承担配置与分析工作，负责参数配置校验逻辑、SQLite数据库设计与读写功能、统计指标计算模块以及历史记录管理功能的编码实现。',
        'problems': [
            ['1', 'SQLite数据库写入性能瓶颈', '仿真过程中每个事件步都产生一条快照记录，逐条写入在高到达率配置下频繁I/O严重拖慢仿真速度。', '设计快照缓冲区机制，每积累50条记录批量写入一次，将写入性能提升约50倍。'],
            ['2', '座位利用率计算超过100%', '初始实现中有效仿真时间仅取配置时长，但部分学生在配置时长后继续就餐。', '将有效仿真时间修正为max(配置时长, 实际结束时间)，确保利用率在0-100%范围内。'],
            ['3', '时间线采样精度与数据量平衡', '按每个事件采样数据量过大，按大间隔又丢失峰值。', '采用按分钟采样策略，没有事件的分钟延用上一分钟数值，兼顾精度和传输效率。'],
        ],
        'task_overview': '本人在本阶段的角色定位为配置与分析，负责数据层和分析层的开发。具体任务如下表所示：',
        'tasks': [
            ['1', 'SQLite数据库表设计与建表', '已完成', '字段类型选择', '参考设计文档'],
            ['2', '参数配置校验函数', '已完成', '边界条件', '逐字段校验+异常捕获'],
            ['3', '快照缓冲区与批量写入', '已完成', '性能瓶颈', '缓冲区批量写入'],
            ['4', '统计指标计算', '已完成', '座位利用率超100%', '修正有效仿真时间'],
            ['5', '历史记录查询接口', '已完成', 'JSON序列化', 'json.loads解析'],
            ['6', '时间线聚合', '已完成', '采样精度选择', '按分钟采样策略'],
        ],
        'task_details': [
            '快照缓冲区机制的设计与实现：最初每个事件步直接写入数据库，在高到达率下仿真速度非常慢。经分析发现瓶颈在频繁的SQLite磁盘I/O。设计了内存缓冲区，积累50条后通过executemany批量写入，性能提升约50倍。这个过程让我理解了缓冲区和批量操作在数据库编程中的重要性。',
            '统计指标计算模块的完善：get_statistics()需要计算6项统计指标和2条时间线数据。最大挑战是座位利用率——需要同时考虑所有学生的就餐时长总和与有效仿真时间的关系。通过数学推导和测试验证，确保了计算结果的正确性。',
        ],
        'tech_gains': [
            'SQLite数据库编程：掌握了建表、批量插入（executemany）、聚合查询（子查询、COUNT、MAX）等操作；',
            '数据缓冲与批量写入策略：理解了缓冲区在数据库性能优化中的应用；',
            '统计指标设计与计算：学会了如何设计有实际意义的仿真统计指标并准确实现；',
            '数据序列化与反序列化：掌握了JSON在前后端数据交换和数据库存储中的使用；',
            '时间序列数据处理：学会了按时间窗口聚合采样的数据处理方法。',
        ],
        'reflection': '本阶段最大的收获是理解了"数据是系统的血液"。一个小小的计算公式错误（如座位利用率超100%）就会直接影响用户对系统的信任度。这让我意识到数据处理代码需要格外注重正确性，每个公式都应该有对应的测试用例来保障。',
        'next_plans': [
            '参与系统集成联调测试，重点验证数据从仿真引擎到数据库再到前端图表的完整链路；',
            '对历史记录查询性能进行优化，考虑为snapshot表添加索引；',
            '配合团队完成联调测试报告的撰写；',
            '根据联调结果优化统计指标的展示效果。',
        ],
    })


# ═══════════════════════════════════════════════════════
if __name__ == '__main__':
    print('正在生成开发阶段交付物（按课程格式规范）…')
    gen_communication_record()
    gen_task_division()
    gen_all_test_reports()
    gen_all_dev_reports()
    print(f'\n全部完成！文件保存在：{OUT}')
    print('\n提示：在 Word 中打开文档后，右键目录区域 → "更新域" → "更新整个目录" 即可生成目录。')
